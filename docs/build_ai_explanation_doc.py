import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "docs" / "Penjelasan_Kode_AI_Keluhan_Perbankan.docx"
METRICS_PATH = ROOT / "reports" / "metrics.json"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def add_heading(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level < 3 else 120)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_doc():
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8"))
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Penjelasan Kode AI Klasifikasi Keluhan Pelanggan Perbankan")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Topik 10 UAS Artificial Intelligence - Naive Bayes vs Logistic Regression").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    add_heading(doc, "Ringkasan Keputusan", 1)
    add_body(
        doc,
        "Project ini dibuat sebagai aplikasi AI end-to-end untuk mengklasifikasikan keluhan pelanggan perbankan. "
        "Saya memilih pendekatan NLP klasik karena sesuai dengan topik 10: teks keluhan diubah menjadi fitur TF-IDF, "
        "lalu dievaluasi menggunakan dua model supervised learning, yaitu Naive Bayes dan Logistic Regression.",
    )
    add_bullets(
        doc,
        [
            "Dataset CFPB dipetakan ulang menjadi 8 kategori bisnis perbankan yang lebih jelas: account_management, account_closure, card_dispute, credit_reporting, fees_interest, fraud_scam, loan_mortgage, dan payment_transfer.",
            "TF-IDF dipakai karena ringan, mudah dijelaskan, dan cocok untuk baseline klasifikasi teks.",
            "Naive Bayes dipakai sebagai baseline cepat untuk data teks.",
            "Logistic Regression dipakai sebagai pembanding linear multiclass yang biasanya kuat untuk fitur sparse seperti TF-IDF.",
            "Web AI dibuat dalam dua bentuk: Streamlit sesuai rubrik dan server Python bawaan agar demo tetap bisa berjalan tanpa instalasi tambahan.",
        ],
    )

    add_heading(doc, "Arsitektur Sistem", 1)
    add_body(doc, "Alur sistem mengikuti pola input layer, preprocessing, inference, lalu output JSON.")
    add_numbered(
        doc,
        [
            "Pengguna memasukkan teks keluhan melalui web UI, Streamlit, CLI, atau API.",
            "Teks dibersihkan melalui case folding, pembersihan karakter, stopword removal, dan stemming.",
            "Teks hasil preprocessing diubah menjadi vektor numerik dengan TF-IDF.",
            "Model Naive Bayes atau Logistic Regression menghasilkan probabilitas setiap kelas.",
            "Sistem mengembalikan kategori utama, confidence score, top 3 prediksi, dan response JSON.",
        ],
    )
    add_body(
        doc,
        "Diagram arsitektur: Input Layer -> Preprocessing NLP -> TF-IDF Vectorizer -> Model Inference -> Output JSON/Web Result.",
    )

    add_heading(doc, "Struktur File Kode", 1)
    add_table(
        doc,
        ["Path", "Fungsi"],
        [
            ["data/complaints.csv.zip", "Dataset mentah CFPB Consumer Complaint Database."],
            ["data/processed_banking_complaints.csv", "Subset olahan hasil filter narasi, pemetaan kategori bisnis, dan balancing per label."],
            ["src/preprocessing.py", "Normalisasi teks, tokenisasi, stopword removal, dan stemming."],
            ["src/vectorizer.py", "Implementasi TF-IDF sederhana berbasis NumPy."],
            ["src/models.py", "Implementasi Multinomial Naive Bayes dan Logistic Regression multiclass."],
            ["src/train.py", "Training, split data, evaluasi, dan penyimpanan model."],
            ["src/predict.py", "Fungsi inference dan format response JSON."],
            ["src/recommendations.py", "Rule-based expert system untuk solusi awal, prioritas, SLA, dan divisi tujuan."],
            ["web_app.py", "Web AI lokal dengan endpoint POST /predict."],
            ["app.py", "Aplikasi Streamlit untuk demo sesuai rubrik."],
            ["reports/metrics.json", "Hasil evaluasi model."],
        ],
    )

    add_heading(doc, "Alasan Pemilihan Komponen", 1)
    add_heading(doc, "Integrasi Dua Teknik AI", 2)
    add_body(
        doc,
        "Project ini tidak berhenti pada klasifikasi. Teknik pertama adalah NLP dan supervised learning untuk memprediksi topik keluhan. "
        "Teknik kedua adalah rule-based expert system yang menggunakan label hasil model sebagai input untuk menentukan solusi awal, prioritas, SLA, dan unit penanganan. "
        "Integrasi ini membuat output lebih operasional karena sistem bukan hanya mengatakan masalahnya apa, tetapi juga menyarankan tindakan awal.",
    )

    add_heading(doc, "Preprocessing Bahasa Indonesia", 2)
    add_body(
        doc,
        "Preprocessing diperlukan karena teks keluhan biasanya berisi variasi huruf besar, tanda baca, dan kata umum yang tidak membantu klasifikasi. "
        "Karena dataset CFPB berbahasa Inggris, stopword bahasa Inggris juga ditambahkan. Kode tetap mempertahankan dukungan PySastrawi jika nanti "
        "dataset diterjemahkan atau diganti menjadi Bahasa Indonesia.",
    )

    add_heading(doc, "TF-IDF", 2)
    add_body(
        doc,
        "TF-IDF dipilih karena memberi bobot lebih tinggi pada kata yang penting untuk dokumen tertentu tetapi tidak terlalu umum di seluruh dataset. "
        "Contohnya kata OTP, transfer, biaya, deposito, dan fraud dapat menjadi sinyal kuat untuk kategori tertentu.",
    )

    add_heading(doc, "Naive Bayes", 2)
    add_body(
        doc,
        "Naive Bayes dipilih sebagai baseline karena sederhana, cepat, dan lazim untuk klasifikasi dokumen. "
        "Model ini cocok untuk melihat apakah kata-kata kunci dalam keluhan sudah cukup informatif untuk memisahkan topik.",
    )

    add_heading(doc, "Logistic Regression", 2)
    add_body(
        doc,
        "Logistic Regression dipilih sebagai pembanding karena model linear ini mampu mempelajari bobot fitur secara diskriminatif. "
        "Pada data teks berbasis TF-IDF, pendekatan linear sering menjadi baseline kuat sebelum memakai model deep learning yang lebih berat.",
    )

    add_heading(doc, "Hasil Evaluasi", 1)
    model_rows = []
    for model_name, report in metrics["models"].items():
        model_rows.append(
            [
                model_name,
                report["accuracy"],
                report["macro_precision"],
                report["macro_recall"],
                report["macro_f1"],
            ]
        )
    add_table(doc, ["Model", "Accuracy", "Macro Precision", "Macro Recall", "Macro F1"], model_rows)
    add_body(
        doc,
        f"Model terbaik pada baseline ini adalah {metrics['best_model']}. Dataset saat ini berisi "
        f"{metrics['dataset']['total_rows']} baris, dengan {metrics['dataset']['train_rows']} data latih dan "
        f"{metrics['dataset']['test_rows']} data uji. Nilai metrik harus dibaca sebagai baseline awal karena dataset masih kecil.",
    )

    add_heading(doc, "API JSON", 1)
    add_body(doc, "Endpoint utama web lokal adalah POST /predict dengan payload JSON.")
    add_table(
        doc,
        ["Jenis", "Contoh"],
        [
            ["Request", '{"text": "Transfer gagal tetapi saldo terpotong", "model": "naive_bayes"}'],
            ["Response 200", '{"status": 200, "prediction": "transfer", "confidence": 0.82}'],
            ["Recommendation", '{"priority": "High", "target_unit": "Fraud Investigation", "actions": ["Blokir sementara akun berisiko"]}'],
            ["Response 400", '{"status": 400, "error": "Teks keluhan tidak boleh kosong."}'],
        ],
    )

    add_heading(doc, "Contoh Rekomendasi Sistem", 1)
    add_table(
        doc,
        ["Label", "Prioritas", "Unit Tujuan", "Contoh Solusi"],
        [
            ["fraud_scam", "Critical", "Fraud Investigation / Security Team", "Blokir kanal berisiko dan kumpulkan bukti transaksi."],
            ["fees_interest", "Medium", "Billing / Fee Dispute Team", "Cocokkan biaya dengan tarif dan ajukan koreksi jika salah tagih."],
            ["payment_transfer", "High", "Payment Operation / Transaction Support", "Rekonsiliasi transaksi dan pantau reversal dana."],
            ["account_management", "Medium", "Account Support / Customer Service", "Verifikasi identitas dan bantu reset akses akun/PIN."],
        ],
    )

    add_heading(doc, "Edge Case dan Robustness", 1)
    add_bullets(
        doc,
        [
            "Input kosong dikembalikan sebagai status 400 dengan pesan yang mudah dipahami.",
            "Model yang tidak tersedia dikembalikan sebagai status 400 dan sistem menampilkan daftar model valid.",
            "Jika file model belum ada, fungsi prediksi otomatis menjalankan training terlebih dahulu.",
            "Web app menangani JSON tidak valid dengan pesan error tanpa membuat server berhenti.",
        ],
    )

    add_heading(doc, "Keterbatasan dan Pengembangan", 1)
    add_body(
        doc,
        "Dataset saat ini memakai CFPB Consumer Complaint Database dari file complaints.csv.zip. Karena dataset asli sangat besar dan label mentahnya terlalu granular, "
        "training memakai subset narasi yang memiliki teks keluhan lalu memetakan data ke kategori bisnis perbankan yang lebih stabil. Evaluasi dapat ditingkatkan "
        "dengan memperbesar MAX_ROWS_PER_LABEL, memakai cross-validation, analisis error per kelas, dan menambah model pembanding seperti Linear SVM.",
    )

    add_heading(doc, "Referensi Teknis", 1)
    add_bullets(
        doc,
        [
            "scikit-learn TfidfVectorizer: https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
            "scikit-learn MultinomialNB: https://scikit-learn.org/stable/modules/generated/sklearn.naive_bayes.MultinomialNB.html",
            "scikit-learn LogisticRegression: https://scikit-learn.org/stable/modules/generated/sklearn.linear_model.LogisticRegression.html",
            "PySastrawi Indonesian Stemmer: https://github.com/har07/PySastrawi",
            "CFPB Consumer Complaint Database: https://www.consumerfinance.gov/data-research/consumer-complaints/",
        ],
    )

    add_heading(doc, "Cara Menjalankan", 1)
    add_body(doc, "Instalasi dependency Streamlit: pip install -r requirements.txt")
    add_body(doc, "Training model lokal: python -m src.train --print")
    add_body(doc, "Training dari API dataset: set DATA_API_URL=http://127.0.0.1:8765/complaints lalu jalankan python -m src.train --print")
    add_body(doc, "Web AI lokal: python web_app.py lalu buka http://127.0.0.1:8501")
    add_body(doc, "Streamlit: pip install -r requirements.txt lalu streamlit run app.py")

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
